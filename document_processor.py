"""
Document processor for handling multiple file types
"""
import os
from pathlib import Path
from typing import List, Dict, Any
import PyPDF2
import docx
from PIL import Image
import easyocr
import numpy as np
import speech_recognition as sr
from pydub import AudioSegment
from moviepy.editor import VideoFileClip
import tempfile
from llama_index.core import Document


class DocumentProcessor:
    """Handles processing of various document types"""
    
    def __init__(self):
        self.recognizer = sr.Recognizer()
    
    def process_file(self, file_path: str) -> List[Document]:
        """
        Process a file based on its extension and return LlamaIndex Documents
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            List of LlamaIndex Document objects
        """
        extension = Path(file_path).suffix.lower()
        
        processors = {
            '.txt': self._process_text,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image,
            '.gif': self._process_image,
            '.bmp': self._process_image,
            '.mp3': self._process_audio,
            '.wav': self._process_audio,
            '.m4a': self._process_audio,
            '.ogg': self._process_audio,
            '.mp4': self._process_video,
            '.avi': self._process_video,
            '.mov': self._process_video,
            '.mkv': self._process_video,
            '.md': self._process_text,
        }
        
        processor = processors.get(extension)
        if processor:
            return processor(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _process_text(self, file_path: str) -> List[Document]:
        """Process text files"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        return [Document(
            text=text,
            metadata={
                'file_name': Path(file_path).name,
                'file_type': 'text',
                'file_path': file_path
            }
        )]
    
    def _process_pdf(self, file_path: str) -> List[Document]:
        """Process PDF files"""
        documents = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                
                if text.strip():
                    documents.append(Document(
                        text=text,
                        metadata={
                            'file_name': Path(file_path).name,
                            'file_type': 'pdf',
                            'page_number': page_num + 1,
                            'file_path': file_path
                        }
                    ))
        
        return documents
    
    def _process_docx(self, file_path: str) -> List[Document]:
        """Process DOCX files"""
        doc = docx.Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        return [Document(
            text=text,
            metadata={
                'file_name': Path(file_path).name,
                'file_type': 'docx',
                'file_path': file_path
            }
        )]
    
    def _process_image(self, file_path: str) -> List[Document]:
        """Process image files using EasyOCR (OpenCV-based)"""
        try:
            # Initialize EasyOCR reader (lazy load to save memory)
            # using cpu=True for compatibility with free tier cloud instances
            if not hasattr(self, 'reader'):
                self.reader = easyocr.Reader(['en'], gpu=False)
            
            # Read text directly from file path
            result = self.reader.readtext(file_path, detail=0)
            text = " ".join(result)
            
            image = Image.open(file_path)
            
            if not text.strip():
                text = f"[Image file: {Path(file_path).name}. No text detected via OCR.]"
            
            return [Document(
                text=text,
                metadata={
                    'file_name': Path(file_path).name,
                    'file_type': 'image',
                    'file_path': file_path,
                    'image_size': f"{image.width}x{image.height}",
                    'ocr_engine': 'easyocr'
                }
            )]
        except Exception as e:
            return [Document(
                text=f"[Image file: {Path(file_path).name}. Error processing: {str(e)}]",
                metadata={
                    'file_name': Path(file_path).name,
                    'file_type': 'image',
                    'file_path': file_path,
                    'error': str(e)
                }
            )]

    
    def _process_audio(self, file_path: str) -> List[Document]:
        """Process audio files using speech recognition"""
        try:
            # Convert to WAV if needed
            audio = AudioSegment.from_file(file_path)
            
            # Export to temporary WAV file
            with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_wav:
                audio.export(temp_wav.name, format='wav')
                temp_wav_path = temp_wav.name
            
            # Perform speech recognition
            with sr.AudioFile(temp_wav_path) as source:
                audio_data = self.recognizer.record(source)
                try:
                    text = self.recognizer.recognize_google(audio_data)
                except sr.UnknownValueError:
                    text = "[Audio file: Google Speech Recognition could not understand audio]"
                except sr.RequestError as e:
                    text = f"[Audio file: Could not request results; {e}]"
            
            # Clean up temp file
            os.unlink(temp_wav_path)
            
            return [Document(
                text=text,
                metadata={
                    'file_name': Path(file_path).name,
                    'file_type': 'audio',
                    'file_path': file_path,
                    'duration_seconds': len(audio) / 1000
                }
            )]
        except Exception as e:
            return [Document(
                text=f"[Audio file: {Path(file_path).name}. Error processing: {str(e)}]",
                metadata={
                    'file_name': Path(file_path).name,
                    'file_type': 'audio',
                    'file_path': file_path,
                    'error': str(e)
                }
            )]
    
    def _process_video(self, file_path: str) -> List[Document]:
        """Process video files by extracting audio and using speech recognition"""
        try:
            # Extract audio from video
            video = VideoFileClip(file_path)
            
            # Export audio to temporary file
            with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_audio:
                video.audio.write_audiofile(temp_audio.name, logger=None)
                temp_audio_path = temp_audio.name
            
            # Process the audio
            audio_docs = self._process_audio(temp_audio_path)
            
            # Update metadata
            for doc in audio_docs:
                doc.metadata.update({
                    'file_name': Path(file_path).name,
                    'file_type': 'video',
                    'file_path': file_path,
                    'duration_seconds': video.duration,
                    'fps': video.fps
                })
            
            # Clean up
            os.unlink(temp_audio_path)
            video.close()
            
            return audio_docs
        except Exception as e:
            return [Document(
                text=f"[Video file: {Path(file_path).name}. Error processing: {str(e)}]",
                metadata={
                    'file_name': Path(file_path).name,
                    'file_type': 'video',
                    'file_path': file_path,
                    'error': str(e)
                }
            )]


def get_file_type_category(file_path: str) -> str:
    """Determine the category of a file"""
    extension = Path(file_path).suffix.lower()
    
    if extension in ['.txt', '.pdf', '.docx', '.doc', '.md']:
        return 'text'
    elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
        return 'image'
    elif extension in ['.mp3', '.wav', '.m4a', '.ogg']:
        return 'audio'
    elif extension in ['.mp4', '.avi', '.mov', '.mkv']:
        return 'video'
    else:
        return 'unknown'
